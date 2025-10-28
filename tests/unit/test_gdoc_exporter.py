"""Unit tests for i4g.reports.gdoc_exporter."""

import os
from docx import Document

from i4g.reports import gdoc_exporter

def test_export_report_creates_docx_file(tmp_path):
    """Ensure the export_report function creates a valid .docx file."""
    # Override the default reports dir to use a temporary directory
    gdoc_exporter.Path.cwd = lambda: tmp_path

    title = "Sample Report"
    content = "This is a test report."
    result = gdoc_exporter.export_report(title, content)

    # 1. Check the return value
    assert result["mode"] == "docx"
    assert "local_path" in result
    local_path = result["local_path"]

    # 2. Check that the file exists and has the right extension
    assert os.path.exists(local_path)
    assert local_path.endswith(".docx")

    # 3. Check the content of the .docx file
    doc = Document(local_path)
    assert doc.paragraphs[0].text == title
    assert doc.paragraphs[0].style.name == "Heading 1"
    assert doc.paragraphs[1].text == content