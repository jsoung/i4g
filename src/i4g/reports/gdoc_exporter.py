"""Docx Exporter for i4g Reports.

Supports local file export to .docx format.
"""

from __future__ import annotations
import datetime
from pathlib import Path
from typing import Dict, Optional

import docx

def export_report(
    title: str,
    content: str,
) -> Dict[str, Optional[str]]:
    """Export a rendered report to a local .docx file.

    Args:
        title: The report title.
        content: The report content (Markdown or text).

    Returns:
        A dictionary with:
            - "local_path": path to the saved .docx file.
            - "mode": "docx".
    """
    reports_dir = Path("reports")
    reports_dir.mkdir(exist_ok=True)
    filename = f"{title.replace(' ', '_')}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    path = reports_dir / filename

    document = docx.Document()
    document.add_heading(title, level=1)
    document.add_paragraph(content)
    document.save(path)

    return {"local_path": str(path.resolve()), "mode": "docx"}